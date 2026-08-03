"""
extractor.py
------------
Content extraction stage — the second real step of the pipeline (after input.py receives the file and scanner.py confirms it's safe).

Purpose:
Pulls out everything useful from a document: normal text directly where that's trustworthy, and anything image-based or layout-complex (charts, scanned receipts, dense tables) via a call out to the vision model.
Returns one combined result per file.

Design goals:
1. Never throw away good work because of a partial failure. If normal text extraction succeeds but the vision model fails on the images, the caller should still get the normal text back, with a clear note that image analysis didn't complete — not nothing at all.
2. Hard caps on pages/images/text length, enforced HERE, not just trusted to whatever cap the vision model happens to have.
3. Don't blindly trust raw text extraction on every page. Complex layouts (dense tables, multi-column financial statements, heavily-drawn pages) often extract as garbled or scrambled text even though no literal embedded image is present. Per project decision, this extractor errs
   toward routing a page to the vision model whenever ANY unusual layout signal shows up — favoring accuracy over the extra vision-model calls that costs, since the model runs locally rather than as a paid API.
   Known limitation: this page-render heuristic currently only applies to PDFs. DOCX tables are still handled via python-docx's own paragraph/ image extraction — a future improvement would read doc.tables directly (python-docx exposes table cells structurally, no vision model needed
   for those at all), but that's not implemented yet.
4. Match whatever the currently active vision model module actually returns — a structured VisionAnalysisResult (accepted / rejected images), not a plain string.
"""

import io                            # in-memory byte buffers for DOCX embedded images
import logging                       # structured logging instead of print()
import tempfile                      # scratch directories for extracted images, auto-cleaned
from pathlib import Path             # safe, OS-independent path handling
from typing import Any, Dict, List, Optional  # type hints so signatures are self-documenting

import fitz                          # pymupdf — reads PDF text, drawings, and renders pages
from docx import Document            # reads DOCX paragraphs and embedded images
from PIL import Image                # decodes embedded image bytes before re-saving to disk

# ============================================================
# VISION MODEL SWITCH
# Just comment / uncomment the model you want to use. NOTE: every module swapped in here must return the same VisionAnalysisResult shape (see vision_model.py) — every call site below assumes .accepted / .rejected lists, not a plain string.
# ============================================================

# from .vision_model import analyze_images                  
# from .vision_model_florence import analyze_images
from ._05_vision_model_llava import analyze_images                  # ← Currently Active (MiniCPM-V)
# from .vision_model_moondream import analyze_images
# from .vision_model_qwen import analyze_images

# ============================================================

logger = logging.getLogger(__name__)  # module-level logger tagged with this file's name


# ============================================================
# CONFIGURATION
# ============================================================

MAX_PAGES_PER_PDF = 200               # stop reading pages beyond this
MAX_IMAGES_PER_DOCUMENT = 25          # covers both embedded images AND full-page renders combined
MAX_EXTRACTED_TEXT_CHARS = 2_000_000  # ~2MB of text; well beyond any real financial document
MAX_TXT_FILE_SIZE_MB = 10.0           # independent safety net for plain text files

# Page layout-complexity thresholds. Deliberately set LOW/loose — per project decision, false positives (sending an actually-simple page to the vision model) are preferred over false negatives (trusting garbled text from a complex page).

COMPLEX_DRAWING_THRESHOLD = 3         # this many vector lines/rects on a page suggests a drawn table
COMPLEX_TEXT_BLOCK_THRESHOLD = 8      # this many separate text blocks suggests multi-column/table layout
COMPLEX_MAX_AVG_WORDS_PER_LINE = 3    # tables tend to have short, sparse lines (cell values, not sentences)
COMPLEX_MIN_LINE_COUNT = 5            # need at least this many lines before the short-line signal counts
PAGE_RENDER_DPI = 200                 # resolution for full-page renders sent to the vision model


class ExtractionError(Exception):
    """
    Raised when content extraction fails outright (unreadable/corrupt file, unsupported type, etc). NOT raised just because the vision model had trouble with some images — see _run_vision_model for why.
    """
    pass


def extract_content(file_path: str) -> Dict[str, Any]:
    """
    Main function: extract all useful content from a document.

    Returns a dict with:
      text            - normal_text + vision_text combined, ready for downstream use
      normal_text     - text pulled directly from the document (no model involved)
      vision_text     - text extracted via the vision model (images + complex pages)
      images_found    - how many images/page-renders were sent to the vision model
      images_rejected - how many of those the vision model skipped, and why
      file_type       - which branch handled this file
    """
    path = Path(file_path).resolve()  # normalize + follow symlinks, same as scanner.py

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    logger.info(f"Extracting content from: {path.name} ({suffix})")

    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        elif suffix == ".docx":
            return _extract_docx(path)
        elif suffix in [".txt", ".md"]:
            return _extract_txt(path)
        elif suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
            return _extract_image(path)
        else:
            raise ExtractionError(f"Unsupported file type: {suffix}")

    except ExtractionError:
        raise  # already a clean, specific error — don't wrap it again
    except Exception as e:
        # Anything unexpected (corrupt file, library-level failure, etc) — convert to our own exception type so callers only ever need to catch one thing, but log the real cause for debugging.
        logger.error(f"Extraction failed for {path.name}: {e}")
        raise ExtractionError(f"Failed to extract content: {str(e)}")


def _run_vision_model(image_paths: List[str]) -> Dict[str, Any]:
    """
    Shared helper for every extraction path that produced images (embedded pictures or full-page renders). Calls the vision model and turns its VisionAnalysisResult into the plain values the rest of this file needs.

    Isolated in its own try/except: if the vision model fails outright (worker crashed past its retry limit, etc), that must NOT destroy normal_text that was already successfully extracted elsewhere.
    """
    if not image_paths:
        return {"vision_text": "", "images_rejected": 0, "rejection_reasons": []}

    if len(image_paths) > MAX_IMAGES_PER_DOCUMENT:
        logger.warning(
            "Document produced %d images/renders, capping to %d before calling the vision model.",
            len(image_paths), MAX_IMAGES_PER_DOCUMENT,
        )
        image_paths = image_paths[:MAX_IMAGES_PER_DOCUMENT]

    try:
        result = analyze_images(image_paths)  # VisionAnalysisResult: .accepted / .rejected
    except Exception as e:
        logger.error(f"Vision model failed on this document's images: {e}")
        return {
            "vision_text": "",
            "images_rejected": len(image_paths),
            "rejection_reasons": [f"Vision model unavailable: {e}"],
        }

    return {
        "vision_text": result.combined_text(),
        "images_rejected": len(result.rejected),
        "rejection_reasons": [r.reason for r in result.rejected],
    }


def _page_has_complex_layout(page: "fitz.Page") -> bool:
    """
    Heuristic: decide whether a page's raw text extraction is untrustworthy enough to route the whole page to the vision model instead. Tuned to err toward flagging pages as complex — a false positive just costs one extra (local, free) vision-model call; 
    a false negative means silently feeding garbled table data into a financial prediction.
    """
    try:
        # Signal 1: heavily hand-drawn content (lines/rects) is a strong sign of a table rendered as vector graphics rather than real text.
        drawings = page.get_drawings()
        if len(drawings) >= COMPLEX_DRAWING_THRESHOLD:
            return True

        text_dict = page.get_text("dict")
        text_blocks = [b for b in text_dict.get("blocks", []) if b.get("type") == 0]

        # Signal 2: many separate text blocks suggests a multi-column or grid layout rather than flowing paragraphs.
        if len(text_blocks) >= COMPLEX_TEXT_BLOCK_THRESHOLD:
            return True

        # Signal 3: lots of short lines (few words each) is typical of table cell values rather than normal sentences.
        line_word_counts = []
        for block in text_blocks:
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                line_text = "".join(s.get("text", "") for s in spans)
                word_count = len(line_text.split())
                if word_count > 0:
                    line_word_counts.append(word_count)

        if len(line_word_counts) >= COMPLEX_MIN_LINE_COUNT:
            avg_words = sum(line_word_counts) / len(line_word_counts)
            if avg_words <= COMPLEX_MAX_AVG_WORDS_PER_LINE:
                return True

        return False
    except Exception as e:
        # If we can't even analyze the layout, treat it as complex rather than silently trusting text we never actually checked.
        logger.warning(f"Layout complexity check failed, defaulting to vision model: {e}")
        return True


def _render_page_to_image(page: "fitz.Page", temp_dir: Path, page_number: int) -> Optional[str]:
    """
    Renders a full page to a PNG for the vision model. Returns the temp path, or None if rendering fails — a render failure should never take down the whole extraction, just fall back to raw text for that page.
    """
    try:
        pix = page.get_pixmap(dpi=PAGE_RENDER_DPI)
        temp_path = temp_dir / f"page_render_{page_number}.png"
        pix.save(str(temp_path))
        return str(temp_path)
    except Exception as e:
        logger.warning(f"Failed to render page {page_number} for vision model: {e}")
        return None


def _extract_pdf(path: Path) -> Dict[str, Any]:
    doc = fitz.open(path)
    normal_text_parts = []
    image_paths = []

    with tempfile.TemporaryDirectory(prefix="pdf_images_") as temp_dir_str:
        temp_dir = Path(temp_dir_str)
        try:
            for page_number, page in enumerate(doc):
                if page_number >= MAX_PAGES_PER_PDF:
                    logger.warning(
                        "PDF has more than %d pages, stopping early at page %d.",
                        MAX_PAGES_PER_PDF, page_number,
                    )
                    break

                complex_layout = _page_has_complex_layout(page)

                if complex_layout:
                    # Untrustworthy text extraction — prefer a full-page render over the raw (possibly scrambled) text.
                    rendered_path = None
                    if len(image_paths) < MAX_IMAGES_PER_DOCUMENT:
                        rendered_path = _render_page_to_image(page, temp_dir, page_number)
                        if rendered_path:
                            image_paths.append(rendered_path)
                    if rendered_path is None:
                        # Rendering failed or we're over the image cap — fall back to raw text rather than losing the page entirely.
                        text = page.get_text("text").strip()
                        if text:
                            normal_text_parts.append(text)
                else:
                    # Layout looks like normal flowing text — trust it.
                    text = page.get_text("text").strip()
                    if text:
                        normal_text_parts.append(text)

                    # Still pull any literal embedded pictures on an otherwise-simple page (e.g. a logo or a small photo sitting next to normal paragraphs).
                    if len(image_paths) < MAX_IMAGES_PER_DOCUMENT:
                        for img in page.get_images(full=True):
                            if len(image_paths) >= MAX_IMAGES_PER_DOCUMENT:
                                break
                            try:
                                xref = img[0]
                                base_image = doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]

                                temp_path = temp_dir / f"image_{page_number}_{xref}.{image_ext}"
                                temp_path.write_bytes(image_bytes)
                                image_paths.append(str(temp_path))
                            except Exception as ex:
                                logger.warning(f"Skipping unreadable image: {ex}")
        finally:
            doc.close()

        normal_text = _cap_text_length("\n\n".join(normal_text_parts).strip())
        vision = _run_vision_model(image_paths)

    final_text = _combine_text(normal_text, vision["vision_text"])

    return {
        "text": final_text,
        "normal_text": normal_text,
        "vision_text": vision["vision_text"],
        "images_found": len(image_paths),
        "images_rejected": vision["images_rejected"],
        "rejection_reasons": vision["rejection_reasons"],
        "file_type": "pdf",
    }


def _extract_docx(path: Path) -> Dict[str, Any]:
    doc = Document(path)
    normal_text_parts = []
    image_paths = []

    with tempfile.TemporaryDirectory(prefix="docx_images_") as temp_dir:
        for para in doc.paragraphs:
            if para.text.strip():
                normal_text_parts.append(para.text.strip())

        for rel in doc.part.rels.values():
            if len(image_paths) >= MAX_IMAGES_PER_DOCUMENT:
                break
            # Check the actual content type of the related part rather than just string-matching the target path — more reliable, since a target path could contain "image" without actually being one.
            content_type = getattr(rel.target_part, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            try:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data))

                temp_path = Path(temp_dir) / f"image_{len(image_paths)}.png"
                image.save(str(temp_path))
                image_paths.append(str(temp_path))
            except Exception as ex:
                logger.warning(f"Skipping unreadable image in docx: {ex}")

        normal_text = _cap_text_length("\n\n".join(normal_text_parts).strip())
        vision = _run_vision_model(image_paths)

    final_text = _combine_text(normal_text, vision["vision_text"])

    return {
        "text": final_text,
        "normal_text": normal_text,
        "vision_text": vision["vision_text"],
        "images_found": len(image_paths),
        "images_rejected": vision["images_rejected"],
        "rejection_reasons": vision["rejection_reasons"],
        "file_type": "docx",
    }


def _extract_txt(path: Path) -> Dict[str, Any]:
    # Independent size safety net — see the note in the module docstring about the current mismatch between this and scanner.py's allowed types (worth resolving one way or the other).
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_TXT_FILE_SIZE_MB:
        raise ExtractionError(f"Text file too large ({size_mb:.2f}MB). Max: {MAX_TXT_FILE_SIZE_MB}MB")

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = _cap_text_length(f.read().strip())

    return {
        "text": text,
        "normal_text": text,
        "vision_text": "",
        "images_found": 0,
        "images_rejected": 0,
        "rejection_reasons": [],
        "file_type": "txt",
    }


def _extract_image(path: Path) -> Dict[str, Any]:
    vision = _run_vision_model([str(path)])

    return {
        "text": vision["vision_text"],
        "normal_text": "",
        "vision_text": vision["vision_text"],
        "images_found": 1,
        "images_rejected": vision["images_rejected"],
        "rejection_reasons": vision["rejection_reasons"],
        "file_type": "image",
    }


def _cap_text_length(text: str) -> str:
    """
    Truncates extracted text to a sane hard ceiling. Prevents an unusually large document from producing a multi-megabyte string that flows uncapped into the cache DB / model prompt downstream.
    """
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        logger.warning(
            "Extracted text length %d exceeds cap %d, truncating.",
            len(text), MAX_EXTRACTED_TEXT_CHARS,
        )
        return text[:MAX_EXTRACTED_TEXT_CHARS] + "\n\n[TRUNCATED: document exceeded max text length]"
    return text


def _combine_text(normal_text: str, vision_text: str) -> str:
    parts = []
    if normal_text:
        parts.append(normal_text)
    if vision_text:
        parts.append("\n\n--- Content from Images ---\n\n" + vision_text)
    return "\n\n".join(parts).strip()


# ==============================
# Quick Test
# ==============================
if __name__ == "__main__":
    import tempfile

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
    )

    print("=" * 60)
    print("Testing extractor.py...")
    print("=" * 60)

    # Self-contained .txt test — no external sample file needed, just
    # confirms the extractor's own wiring (dispatch by extension, text
    # capping, return shape) works.
    with tempfile.TemporaryDirectory() as tmp_dir:
        test_path = Path(tmp_dir) / "test_sample.txt"
        test_path.write_text("This is a test financial note: total due $123.45.")

        try:
            result = extract_content(str(test_path))
            print(f"✅ Extraction succeeded, file_type={result['file_type']}")
            print(f"   Text: {result['text']}")
        except ExtractionError as e:
            print(f"❌ Extraction failed: {e}")

    # Exercising the PDF/DOCX/image branches (and therefore the vision
    # model) needs a real sample file — change this path and uncomment:
    # real_test_path = "/path/to/your/sample.pdf"
    # print(extract_content(real_test_path))